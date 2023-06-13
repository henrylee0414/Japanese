#導入所需函式庫
import requests
import random
import time
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


#迴圈從5倒數到3
for num in range(5, 2, -1):
    
    #建立新的word檔
    doc = Document()
    
    #用requests向網址發送get請求，並使用BeautifulSoup解析
    url = f"https://www.sigure.tw/learn-japanese/vocabulary/n{num}/"
    response = requests.get(url)
    soup = BeautifulSoup(response.text,"html.parser")
    
    #找到網頁中放單字的網址，再次送出請求
    for link in soup.find(id="myTable").find_all('a'):
        response2 = requests.get("https://www.sigure.tw"+link.get("href"))
        soup2 = BeautifulSoup(response2.text,"html.parser")

        # 找到所有需要處理的內容（例如段落、標題和表格）
        contents = soup2.find("div",itemprop="articleBody").find_all(['p', 'h2', 'h3', 'table', 'div'])
    
        #h1先處理
        h1 = soup2.find("h1")
        doc.add_heading(h1.text, level=1)
    
        # 遍歷內容並處理
        for content in contents:
        
            #文字
            if content.name == 'p':
                doc.add_paragraph(content.text)
            
            #標題
            elif content.name in ['h2', 'h3']:
                doc.add_heading(content.text, level=int(content.name[1]))
        
            #note,block
            elif content.name == 'div' and ('note' in content.get('class', []) or 'block' in content.get('class', [])):
                if 'note' in content.get('class', []):
                    doc.add_paragraph(content.text)
                if 'block' in content.get('class', []):
                    doc.add_paragraph(content.text)

            #表格
            elif content.name == 'table':
                # 獲取表格的所有行和列
                rows = content.find_all('tr')
                num_cols = len(rows[0].find_all(['th', 'td']))

                # 創建一個新的 Word 表格
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = 'Table Grid'

                # 遍歷表格的行和列
                for row in rows:
                    table_row = table.add_row().cells
                    cells = row.find_all(['th', 'td'])

                    if len(cells) == num_cols:
                        # 將單元格內容添加到 Word 表格中
                        for i, cell in enumerate(cells):
                            table_row[i].text = cell.text

                            # 設置邊框樣式
                            borders = table_row[i]._element.xpath('.//w:tcBorders')
                            if borders:
                                borders = borders[0]
                            else:
                                tc_pr = table_row[i]._element.get_or_add_tcPr()
                                borders = OxmlElement('w:tcBorders')
                                tc_pr.append(borders)

                            # 設置邊框寬度和顏色
                            for border in borders:
                                for attr in border:
                                    attr.attrib['sz'] = '1'  # 設置邊框寬度
                                    attr.attrib['color'] = 'auto'  # 設置邊框顏色

                            # 設置垂直對齊方式
                            table_row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        print("抓取成功")
        
        #為下一章設置下一頁      
        doc.add_page_break()
        
        #設置時間delay
        delay_choices = [1, 5, 3, 6, 2]
        delay = random.choice(delay_choices)
        print('此次延遲時間為',delay,'秒')
        time.sleep(delay)

        #保存文件
        doc.save(f'N{num}單字.docx')
        print('已儲存至',f'N{num}單字.docx\n')